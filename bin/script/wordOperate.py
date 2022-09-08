# This Python file uses the following encoding: utf-8

import json
import os
from docx import Document   #用来建立一个word对象
from docx.shared import Pt  #用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  #设置字体
from docx.shared import RGBColor  #设置字体的颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH  #设置对其方式

def callTest(string):
    print('传入的字符串为：' + string)
    #返回结果
    reStr = 'Python返回值'
    return reStr

# 根据JSON内容生成 word文档
def generateWord(strContent):
    print('call generateWord')
    jsonData = json.loads(strContent)
    #print("input data:",jsonData)

    # 全局字体信息
    globalStyleDictionary = {
        'fontName': u'宋体',
        'fontSize': 10.5,
        'fontColor': RGBColor(0,0,0),
        'space_before': 15,     # 段前
        'space_after': 15       # 断后间距
    }

    document = Document()
    document.styles['Normal'].font.name = globalStyleDictionary['fontName']
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), globalStyleDictionary['fontName'])
    document.styles['Normal'].font.size = Pt(globalStyleDictionary['fontSize'])
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)

    savePath = jsonData['savePath']
    line_spacing = jsonData['line_spacing']
    contentArray = jsonData['content']

    header = document.sections[0].header
    footer = document.sections[0].footer
    headerGraph = header.add_paragraph(jsonData['header'])
    footerGraph = footer.add_paragraph(jsonData['footer'])
    headerGraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footerGraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for itemData in contentArray:
        print('item:', itemData['text'])

        if itemData['type'] == '0':
            # 标题处理
            color = globalStyleDictionary['fontColor']
            if ('color' in itemData.keys()):
                color = hex_to_RGB(itemData['color'])

            para_heading = document.add_heading('', level=itemData['level'])
            headingRun = para_heading.add_run(itemData['text'])
            para_heading.alignment = converAlignment(itemData['alignment'])
            headingRun.bold = itemData['bold']
            headingRun.italic = itemData['italic']
            headingRun.font.name = globalStyleDictionary['fontName']
            headingRun.font.strike = itemData['strike']
            headingRun._element.rPr.rFonts.set(qn('w:eastAsia'), globalStyleDictionary['fontName'])
            headingRun.font.color.rgb = color
        elif itemData['type'] == '1':
            # 普通文本处理
            #document.add_paragraph('')
            paragraph = document.add_paragraph(itemData['text'])

            if('line_spacing' in itemData.keys()):
                line_spacing = itemData['line_spacing']

            # 下面的代码是为了首行缩进 2 字符
            paragraph.style.font.size = Pt(globalStyleDictionary['fontSize'])
            paragraph.paragraph_format.first_line_indent = paragraph.style.font.size * 2

            # 设置段前段后间距信息
            space_before = globalStyleDictionary['space_before']
            space_after = globalStyleDictionary['space_after']

            if('space_before' in itemData.keys()):
                space_before = itemData['space_before']
            if('space_after' in itemData.keys()):
                space_after = itemData['space_after']

            paragraph.paragraph_format.line_spacing = line_spacing
            paragraph.paragraph_format.space_before  = Pt(space_before)  # 段前
            paragraph.paragraph_format.space_after = Pt(space_after)     # 段后
        elif itemData['type'] == '2':
            # 图片处理
            document.add_paragraph('')
            imageObj = document.add_picture(itemData['picture'],width=Inches(5), height=Inches(5))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif itemData['type'] == '3':
            # 表格处理,默认居中对齐
            nRows = itemData['rows']
            nColumns = itemData['columns']
            mergeCells = itemData['mergeCells']

            table = document.add_table(rows=nRows, cols=nColumns, style ='Table Grid')
            table.style.font.size = Pt(globalStyleDictionary['fontSize'])
            table.style.font.name = globalStyleDictionary['fontName']
            table.style.font.color.rgb = globalStyleDictionary['fontColor']
            table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 处理合并单元格场景
            for index ,item in enumerate(mergeCells):
                mergeBegin = item['begin']
                mergeEnd = item['end']
                if (isValidTableIndex(mergeBegin, nRows, nColumns) and isValidTableIndex(mergeEnd, nRows, nColumns) ):
                    table.cell(mergeBegin[0], mergeBegin[1]).merge(table.cell(mergeEnd[0],mergeEnd[1]))

            tableCellArray = itemData['tableCell']
            for r in range(itemData['rows']):
                for c in range(itemData['columns']):
                    cell = table.cell(r, c)
                    itemIndex = r*nColumns + c
                    #print('itemIndex:',r,c, itemIndex)
                    if (itemIndex > len(tableCellArray)-1):
                        continue

                    cell.text = tableCellArray[itemIndex]['text']

                    cellGraph = cell.paragraphs[0]
                    cellGraph.alignment = converAlignment(tableCellArray[itemIndex]['alignment'])
                    cellGraph.add_run('')
                    cellRun = cellGraph.runs[0]                   
                    cellRun.font.bold = tableCellArray[itemIndex]['bold']
                    cellRun.font.italic = tableCellArray[itemIndex]['italic']
                    cellRun.font.color.rgb = hex_to_RGB(tableCellArray[itemIndex]['color'])
        else:
            document.add_paragraph('')
            paragraph = document.add_paragraph(itemData['text'])
    document.save(savePath)

    if jsonData['openFile']:
        os.startfile(savePath)

    return True

# 对齐方式类型转换
def converAlignment(strType):
    if (strType == 'left'):
        return WD_ALIGN_PARAGRAPH.LEFT
    elif (strType == 'center'):
        return WD_ALIGN_PARAGRAPH.CENTER
    elif (strType == 'right'):
        return WD_ALIGN_PARAGRAPH.RIGHT
    elif (strType == 'justify'):
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


# 16进制颜色格式颜色转换为RGB格式
def hex_to_RGB(hex):
    r = int(hex[1:3],16)
    g = int(hex[3:5],16)
    b = int(hex[5:7], 16)
    return RGBColor(r, g, b)

# 表格单元格索引非法判断
def isValidTableIndex(cell, rows, columns):
    if (len(cell) != 2):
        return False
    if (cell[0] < 0 or cell[0] > rows):
        return False
    if (cell[1] < 0 or cell[1] > columns):
        return False
    return True

# 测试
#inputFile = open('./../test.json', encoding='utf-8')
#exportData = json.loads(inputFile.read())
#generateWord(json.dumps(exportData))
